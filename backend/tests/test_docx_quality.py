"""Docx rendering quality: structure recognition, flag-marker survival,
disclaimer footer, and faithful render -> extract round-trip."""
from __future__ import annotations

import io

from docx import Document

from models.schema import SLP_DISCLAIMER
from services.docx_html import docx_to_html
from services.docx_renderer import extract_text, render_docx


def _style_names(docx_bytes: bytes) -> dict[str, str]:
    """Map each non-empty paragraph's text -> its style name."""
    document = Document(io.BytesIO(docx_bytes))
    return {p.text.strip(): (p.style.name or "") for p in document.paragraphs if p.text.strip()}


class TestStructureRecognition:
    def test_numbered_clauses_become_headings(self):
        text = (
            "1. Objeto del acuerdo\n"
            "El presente acuerdo regula la inversión.\n"
            "1.1 Definiciones aplicables\n"
            "Los términos en mayúscula tienen el significado asignado."
        )
        styles = _style_names(render_docx(text))
        assert styles["1. Objeto del acuerdo"] == "Heading 1"
        assert styles["1.1 Definiciones aplicables"] == "Heading 2"
        # Body lines stay body paragraphs.
        assert styles["El presente acuerdo regula la inversión."] == "Normal"

    def test_numbered_clauses_render_as_h_tags(self):
        html = docx_to_html(render_docx("1. Objeto\nTexto del cuerpo."))["html"]
        assert "<h1>1. Objeto</h1>" in html
        assert "<p>Texto del cuerpo.</p>" in html

    def test_ordinal_and_article_headers_become_headings(self):
        text = (
            "PRIMERA.- Objeto\n"
            "Cuerpo de la cláusula primera.\n"
            "ARTÍCULO 5 Duración\n"
            "Cuerpo del artículo quinto."
        )
        styles = _style_names(render_docx(text))
        assert styles["PRIMERA.- Objeto"] == "Heading 2"
        assert styles["ARTÍCULO 5 Duración"] == "Heading 2"

    def test_ordinal_word_in_prose_is_not_a_heading(self):
        # "Segundo acuerdo..." is prose, NOT an ordinal header (no separator).
        styles = _style_names(render_docx("Segundo acuerdo del consejo de administración."))
        assert styles["Segundo acuerdo del consejo de administración."] == "Normal"

    def test_all_caps_title_becomes_heading(self):
        styles = _style_names(render_docx("ACTA DE REUNIÓN DEL CONSEJO\nPrimer acuerdo."))
        assert styles["ACTA DE REUNIÓN DEL CONSEJO"] == "Heading 2"

    def test_bullet_and_dash_lists_become_list_items(self):
        text = "Acuerdos:\n- Primer punto\n• Segundo punto"
        styles = _style_names(render_docx(text))
        assert styles["Primer punto"] == "List Bullet"
        assert styles["Segundo punto"] == "List Bullet"

    def test_monetary_amount_not_mistaken_for_clause(self):
        styles = _style_names(render_docx("1.000.000 de euros se distribuirán a los inversores."))
        assert styles["1.000.000 de euros se distribuirán a los inversores."] == "Normal"


class TestFlagMarkers:
    def test_missing_and_deviation_markers_survive_extract_verbatim(self):
        text = (
            "Cláusula con [MISSING: fecha de cierre] pendiente.\n"
            "Otra cláusula [DEVIATION: se aparta del precedente por petición del cliente].\n"
            "[REFINEMENT-UNCLEAR: instrucción contradictoria]"
        )
        extracted = extract_text(render_docx(text))
        assert "[MISSING: fecha de cierre]" in extracted
        assert "[DEVIATION: se aparta del precedente por petición del cliente]" in extracted
        assert "[REFINEMENT-UNCLEAR: instrucción contradictoria]" in extracted

    def test_markers_rendered_with_distinct_styling(self):
        # The marker run is bold + coloured; surrounding text is a plain run.
        document = Document(io.BytesIO(render_docx("Texto con [MISSING: importe] al final.")))
        paragraph = next(p for p in document.paragraphs if "MISSING" in p.text)
        marker_run = next(r for r in paragraph.runs if "[MISSING:" in r.text)
        plain_run = next(r for r in paragraph.runs if "Texto con" in r.text)
        assert marker_run.bold is True
        assert marker_run.font.color.rgb is not None
        assert plain_run.bold is not True


class TestDisclaimerFooter:
    def test_disclaimer_rendered_once(self):
        body = "ACTA DE REUNIÓN\nPrimer acuerdo."
        text = f"{body}\n\n{SLP_DISCLAIMER}"
        extracted = extract_text(render_docx(text))
        assert extracted.count(SLP_DISCLAIMER) == 1

    def test_disclaimer_styled_distinctly(self):
        text = f"Cuerpo del documento.\n\n{SLP_DISCLAIMER}"
        document = Document(io.BytesIO(render_docx(text)))
        disclaimer_para = next(p for p in document.paragraphs if SLP_DISCLAIMER in p.text)
        run = disclaimer_para.runs[0]
        assert run.italic is True

    def test_text_after_disclaimer_is_preserved(self):
        # The Level-3 warning is appended AFTER the disclaimer in the pipeline.
        warning = "Este documento se ha generado sin precedente de referencia."
        text = f"Cuerpo.\n\n{SLP_DISCLAIMER}\n\n{warning}"
        extracted = extract_text(render_docx(text))
        assert warning in extracted
        assert SLP_DISCLAIMER in extracted


class TestRoundTrip:
    def test_render_extract_preserves_body_text(self):
        text = (
            "ACTA DE REUNIÓN DEL CONSEJO\n"
            "1. Convocatoria\n"
            "Se convoca a los consejeros conforme a los estatutos.\n"
            "- Punto uno del orden del día\n"
            "El consejo aprueba por unanimidad la propuesta presentada."
        )
        extracted = extract_text(render_docx(text))
        # Body sentences survive verbatim (list markers are absorbed into styles).
        for line in (
            "ACTA DE REUNIÓN DEL CONSEJO",
            "1. Convocatoria",
            "Se convoca a los consejeros conforme a los estatutos.",
            "Punto uno del orden del día",
            "El consejo aprueba por unanimidad la propuesta presentada.",
        ):
            assert line in extracted
