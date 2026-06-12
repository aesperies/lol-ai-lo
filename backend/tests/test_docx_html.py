"""In-browser document viewer: docx -> safe HTML conversion + HTML endpoint."""
from __future__ import annotations

import io
import re

from docx import Document

from services.docx_html import docx_to_html
from services.docx_renderer import render_docx
from services.redline import build_redline
from tests.conftest import auth, seed_precedent

PRECEDENT = (
    "ACTA DE REUNIÓN DEL CONSEJO\n"
    "Se aprueba la llamada de capital por importe de 100.000 euros.\n"
    "La reunión se celebra en Madrid.\n"
    "Cláusula que será eliminada en la nueva versión."
)
GENERATED = (
    "ACTA DE REUNIÓN DEL CONSEJO\n"
    "Se aprueba la llamada de capital por importe de 500.000 euros.\n"
    "La reunión se celebra en Madrid.\n"
    "Cláusula nueva añadida sobre el período de inversión."
)


class TestDraftToHtml:
    def test_renders_headings_bold_and_paragraphs(self):
        # render_docx marks short ALL-CAPS lines as bold runs -> heading heuristic.
        result = docx_to_html(
            render_docx("ACTA DE REUNIÓN DEL CONSEJO\nPrimer acuerdo del consejo.\nSegundo acuerdo del consejo.")
        )
        html = result["html"]
        assert "<h2>ACTA DE REUNIÓN DEL CONSEJO</h2>" in html
        assert "<p>Primer acuerdo del consejo.</p>" in html
        assert "<p>Segundo acuerdo del consejo.</p>" in html
        assert result["stats"] == {"insertions": 0, "deletions": 0}

    def test_bold_italic_underline_runs(self):
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("normal ")
        paragraph.add_run("negrita").bold = True
        paragraph.add_run(" cursiva").italic = True
        paragraph.add_run(" subrayado").underline = True
        buffer = io.BytesIO()
        document.save(buffer)

        html = docx_to_html(buffer.getvalue())["html"]
        assert "<strong>negrita</strong>" in html
        assert "<em> cursiva</em>" in html
        assert "<u> subrayado</u>" in html

    def test_simple_table_rendered(self):
        document = Document()
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Inversor"
        table.cell(0, 1).text = "Importe"
        buffer = io.BytesIO()
        document.save(buffer)

        html = docx_to_html(buffer.getvalue())["html"]
        assert '<table class="doc-table"><tr><td>Inversor</td><td>Importe</td></tr></table>' in html

    def test_text_content_is_escaped(self):
        payload = "Cláusula <script>alert(1)</script> maliciosa & peligrosa."
        html = docx_to_html(render_docx(payload))["html"]
        assert "<script>" not in html
        assert "&lt;script&gt;alert(1)&lt;/script&gt;" in html
        assert "&amp; peligrosa" in html

    def test_only_whitelisted_tags_emitted(self):
        html = docx_to_html(build_redline(PRECEDENT, GENERATED))["html"]
        allowed = {"p", "h1", "h2", "h3", "strong", "em", "u", "ins", "del", "table", "tr", "td", "ul", "ol", "li", "br"}
        assert set(re.findall(r"</?([a-z0-9]+)", html)) <= allowed


class TestRedlineToHtml:
    def test_ins_and_del_marks_with_correct_counts(self):
        redline_bytes = build_redline(PRECEDENT, GENERATED)
        result = docx_to_html(redline_bytes)
        html = result["html"]
        assert '<ins class="rl-ins">' in html
        assert '<del class="rl-del">' in html
        # Counts mirror the revision elements written by services/redline.py.
        xml = Document(io.BytesIO(redline_bytes)).element.xml
        assert result["stats"]["insertions"] == xml.count("<w:ins ")
        assert result["stats"]["deletions"] == xml.count("<w:del ")
        assert result["stats"]["insertions"] > 0
        assert result["stats"]["deletions"] > 0
        # Old amount deleted, new amount inserted.
        assert "100.000" in html and "500.000" in html

    def test_unchanged_text_not_marked(self):
        html = docx_to_html(build_redline(PRECEDENT, GENERATED))["html"]
        madrid = next(b for b in html.split("\n") if "Madrid" in b)
        assert "rl-ins" not in madrid and "rl-del" not in madrid

    def test_redline_text_is_escaped(self):
        redline_bytes = build_redline(
            "Texto antiguo del precedente.", "Texto <script>alert(1)</script> nuevo."
        )
        html = docx_to_html(redline_bytes)["html"]
        assert "<script>" not in html
        assert "&lt;script&gt;alert(1)&lt;/script&gt;" in html


class TestHtmlEndpoint:
    def test_draft_and_redline_html_returned_and_audited(self, wf, client, seed, db):
        seed_precedent(db, gestora_id=seed["gestora_a"]["id"], text="PRECEDENTE ALFA")
        request_id, _ = wf.to_review_pending()

        draft = client.get(
            f"/api/requests/{request_id}/documents/draft/html", headers=auth(seed["client_a"])
        )
        assert draft.status_code == 200, draft.text
        assert "<p>" in draft.json()["html"]
        assert draft.json()["stats"] == {"insertions": 0, "deletions": 0}

        redline = client.get(
            f"/api/requests/{request_id}/documents/redline/html", headers=auth(seed["client_a"])
        )
        assert redline.status_code == 200, redline.text
        assert redline.json()["stats"]["insertions"] > 0

        # Audited with the existing download actions + inline_view metadata.
        entries = [
            e
            for e in db.select("audit_log", action="draft_downloaded")
            if (e.get("metadata") or {}).get("request_id") == request_id
        ]
        assert entries and entries[-1]["metadata"]["mode"] == "inline_view"
        redline_entries = [
            e
            for e in db.select("audit_log", action="redline_downloaded")
            if (e.get("metadata") or {}).get("request_id") == request_id
        ]
        assert redline_entries and redline_entries[-1]["metadata"]["mode"] == "inline_view"

    def test_cross_gestora_access_blocked_with_404(self, wf, client, seed, db):
        seed_precedent(db, gestora_id=seed["gestora_a"]["id"], text="PRECEDENTE ALFA")
        request_id, _ = wf.to_review_pending()
        url = f"/api/requests/{request_id}/documents/draft/html"
        assert client.get(url, headers=auth(seed["client_b"])).status_code == 404
        assert client.get(url, headers=auth(seed["client_a"])).status_code == 200
        # Counsel is cross-gestora by design.
        assert client.get(url, headers=auth(seed["counsel"])).status_code == 200

    def test_counsel_edit_is_internal_only(self, wf, client, seed):
        request_id, _ = wf.to_review_pending()
        response = client.get(
            f"/api/requests/{request_id}/documents/counsel_edit/html",
            headers=auth(seed["client_a"]),
        )
        assert response.status_code == 404

    def test_unknown_version_type_rejected(self, wf, client, seed):
        request_id, _ = wf.to_review_pending()
        response = client.get(
            f"/api/requests/{request_id}/documents/bogus/html", headers=auth(seed["client_a"])
        )
        assert response.status_code == 422

    def test_final_html_requires_validated_status(self, wf, client, seed):
        request_id, _ = wf.to_review_pending()
        response = client.get(
            f"/api/requests/{request_id}/documents/final/html", headers=auth(seed["client_a"])
        )
        assert response.status_code == 409
