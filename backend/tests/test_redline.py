"""Redline engine tests: tracked changes, author, formatting-change exclusion."""
from __future__ import annotations

import io
import re

from docx import Document

from services.docx_renderer import extract_text, render_docx
from services.redline import REDLINE_AUTHOR, build_redline


def _xml(docx_bytes: bytes) -> str:
    return Document(io.BytesIO(docx_bytes)).element.xml


PRECEDENT = (
    "ACTA DE REUNIÓN DEL CONSEJO\n"
    "Se aprueba la llamada de capital por importe de 100.000 euros.\n"
    "La reunión se celebra en Madrid.\n"
    "Cláusula que será eliminada en la nueva versión."
)
GENERATED = (
    "ACTA DE REUNIÓN DEL CONSEJO\n"
    "Se aprueba la llamada de capital por importe de 500.000 euros.\n"
    "La reunión se celebra en Madrid.\n"
    "Cláusula nueva añadida sobre el período de inversión."
)


class TestRedline:
    def test_contains_tracked_insertions_and_deletions(self):
        xml = _xml(build_redline(PRECEDENT, GENERATED))
        assert "<w:ins " in xml
        assert "<w:del " in xml
        # Old amount deleted, new amount inserted.
        assert "100.000" in xml and "500.000" in xml

    def test_author_is_always_lolailo_ai(self):
        xml = _xml(build_redline(PRECEDENT, GENERATED))
        authors = set(re.findall(r'w:author="([^"]+)"', xml))
        assert authors == {REDLINE_AUTHOR}
        assert REDLINE_AUTHOR == "Lol-AI-lo AI"

    def test_unchanged_paragraphs_not_marked(self):
        redline = build_redline(PRECEDENT, GENERATED)
        document = Document(io.BytesIO(redline))
        # "La reunión se celebra en Madrid." is identical -> its paragraph
        # must contain no revision elements.
        for paragraph in document.paragraphs:
            if "Madrid" in paragraph.text:
                paragraph_xml = paragraph._p.xml
                assert "<w:ins " not in paragraph_xml
                assert "<w:del " not in paragraph_xml
                break
        else:
            raise AssertionError("Unchanged paragraph missing from redline")

    def test_word_level_diff_only_marks_changed_tokens(self):
        from docx.oxml.ns import qn

        # The shared prefix of the modified sentence stays unmarked: it must
        # appear in plain w:r runs that hang directly off the paragraph
        # (i.e. not wrapped in w:ins / w:del).
        document = Document(io.BytesIO(build_redline(PRECEDENT, GENERATED)))
        target = next(p for p in document.paragraphs if "llamada de capital" in p._p.xml)
        plain_text = ""
        for child in target._p:
            if child.tag == qn("w:r"):
                t = child.find(qn("w:t"))
                if t is not None and t.text:
                    plain_text += t.text
        assert "llamada de capital" in plain_text
        assert "<w:del " in target._p.xml and "<w:ins " in target._p.xml

    def test_pure_formatting_changes_not_marked(self):
        original = "Se aprueba la llamada de capital.\nLa reunión se celebra en Madrid."
        reformatted = "Se  aprueba   la llamada de capital.\nLa reunión se celebra   en Madrid."
        xml = _xml(build_redline(original, reformatted))
        assert "<w:ins " not in xml
        assert "<w:del " not in xml

    def test_deleted_text_uses_deltext(self):
        xml = _xml(build_redline("Cláusula que desaparece por completo.", "Texto totalmente distinto."))
        assert "<w:delText" in xml

    def test_one_word_change_is_word_level_not_full_paragraph(self):
        from docx.oxml.ns import qn

        precedent = "El plazo del fondo será de diez años desde la fecha de cierre."
        generated = "El plazo del fondo será de doce años desde la fecha de cierre."
        document = Document(io.BytesIO(build_redline(precedent, generated)))
        target = next(p for p in document.paragraphs if "plazo del fondo" in p._p.xml)

        # Only the changed word is marked; the rest stays in plain runs.
        del_text = "".join(
            t.text or ""
            for d in target._p.findall(qn("w:del"))
            for t in d.iter(qn("w:delText"))
        )
        ins_text = "".join(
            t.text or ""
            for d in target._p.findall(qn("w:ins"))
            for t in d.iter(qn("w:t"))
        )
        assert del_text.strip() == "diez"
        assert ins_text.strip() == "doce"
        # Shared words must NOT have been deleted+reinserted wholesale.
        assert "plazo del fondo" not in del_text

    def test_whitespace_only_change_produces_no_revisions(self):
        original = "El consejo aprueba la operación por unanimidad."
        whitespace = "El   consejo  aprueba la   operación por unanimidad."
        xml = _xml(build_redline(original, whitespace))
        assert "<w:ins " not in xml
        assert "<w:del " not in xml

    def test_large_input_falls_back_without_error(self, monkeypatch):
        # Force the large-document path with a tiny threshold so the test stays
        # fast; the coarse paragraph-level diff must still be a valid redline.
        import services.redline as redline_mod

        class _FakeSettings:
            redline_max_paragraphs = 5

        monkeypatch.setattr(redline_mod, "get_settings", lambda: _FakeSettings())

        precedent = "\n".join(f"Cláusula {i} del precedente original." for i in range(50))
        generated = "\n".join(f"Cláusula {i} del documento generado." for i in range(50))
        redline_bytes = build_redline(precedent, generated)

        # Valid .docx, tracked changes present, author preserved.
        xml = _xml(redline_bytes)
        assert "<w:ins " in xml and "<w:del " in xml
        authors = set(re.findall(r'w:author="([^"]+)"', xml))
        assert authors == {REDLINE_AUTHOR}
        # Re-openable as a document.
        assert Document(io.BytesIO(redline_bytes)).paragraphs

    def test_added_and_removed_paragraphs_in_replace_block(self):
        # An old clause with no good match becomes a clean deletion; a brand-new
        # clause becomes a clean insertion (not a forced word-diff pairing).
        precedent = (
            "CLÁUSULA ÚNICA\n"
            "El importe de la inversión asciende a 100.000 euros.\n"
            "Texto antiguo completamente diferente que se elimina."
        )
        generated = (
            "CLÁUSULA ÚNICA\n"
            "El importe de la inversión asciende a 250.000 euros.\n"
            "Una cláusula totalmente nueva sobre gobernanza."
        )
        xml = _xml(build_redline(precedent, generated))
        assert "100.000" in xml and "250.000" in xml
        assert "<w:ins " in xml and "<w:del " in xml


class TestDocxRenderer:
    def test_render_and_extract_roundtrip(self):
        text = "ACTA DE REUNIÓN\nPrimer acuerdo del consejo.\nSegundo acuerdo del consejo."
        extracted = extract_text(render_docx(text))
        for line in text.split("\n"):
            assert line in extracted
