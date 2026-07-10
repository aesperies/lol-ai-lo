"""Render generated document text into a formatted .docx and extract text back.

The model emits plain-text legal documents (paragraphs separated by blank
lines, numbered clauses, ALL-CAPS section titles, dash/bullet lists). This
module maps that structure onto real Word paragraph styles so the output looks
professional and so :mod:`services.docx_html` (which keys off Heading styles +
bold runs) can render a faithful hierarchy.

Two invariants the rest of the platform depends on:

* :func:`extract_text` is the inverse of :func:`render_docx` for *meaningful*
  text — every body line, and crucially every ``[MISSING: ...]`` /
  ``[DEVIATION: ...]`` / ``[REFINEMENT-UNCLEAR: ...]`` flag marker, survives the
  round-trip **verbatim** (RAG bases, redline bases, and the Exit-A ``[MISSING]``
  gate read it back).
* The SLP disclaimer arrives already appended to the text; it is rendered once
  as a clearly-separated footer block and never duplicated.

python-docx only — no new dependencies.
"""
from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from models.schema import SLP_DISCLAIMER

# Base typography / layout.
_BASE_FONT = "Calibri"
_BASE_FONT_SIZE = Pt(11)
_MARGIN_CM = 2.5  # ~1 inch professional margins.

# Inline flag markers the model emits. They MUST survive extract_text verbatim,
# but we highlight them so counsel/client can spot them at a glance.
_FLAG_MARKERS = ("[MISSING:", "[DEVIATION:", "[REFINEMENT-UNCLEAR:")
_FLAG_HIGHLIGHT = RGBColor(0xB0, 0x00, 0x00)  # dark red run colour
# Verifiable-citation marker (grounding Feature 1). A sibling of the flag
# markers but styled distinctly: small, muted, footnote-style (NOT the dark-red
# defect highlight) — a citation, not a defect. Also survives extract_text
# verbatim. The closing-bracket-aware pattern allows the literal ``]`` that may
# appear inside the quoted excerpt (it matches up to the final ``]``).
_SOURCE_MARKER = "[SOURCE:"
_SOURCE_FOOTNOTE = RGBColor(0x66, 0x66, 0x66)  # muted grey run colour
# All inline markers, used to split a line into runs. SOURCE is matched first
# (its bracket may wrap a quote containing ']'); the others are bracket-bounded.
_FLAG_SPLIT_RE = re.compile(
    r'(\[SOURCE:[^\]]*"[^"]*"\s*\]'
    r"|\[(?:MISSING|DEVIATION|REFINEMENT-UNCLEAR):[^\]]*\])"
)

# Structure recognition patterns (applied to a stripped line).
#   "1." / "1.1" / "1.1.1" / "12.3."  -> numbered clause. Groups capped at two
#   digits so monetary amounts ("1.000.000 euros") are NOT mistaken for clauses.
_NUMBERED_CLAUSE_RE = re.compile(r"^(\d{1,2}(?:\.\d{1,2})*)\.?\s+\S")
#   "PRIMERA.-" / "SEGUNDO.-" / "PRIMERA -" / "PRIMERA:" ordinal headers (ES).
#   Must be followed by a header separator (.- . : -) so ordinary prose that
#   merely starts with an ordinal word ("Segundo acuerdo del consejo.") is NOT
#   treated as a heading.
_ORDINAL_HEADER_RE = re.compile(
    r"^(PRIMER[AO]|SEGUND[AO]|TERCER[AO]|CUART[AO]|QUINT[AO]|SEXT[AO]|"
    r"S[ÉE]PTIM[AO]|OCTAV[AO]|NOVEN[AO]|D[ÉE]CIM[AO])\s*(\.-|[.:\-])\s*\S",
    re.IGNORECASE,
)
#   "ARTÍCULO 5" / "ARTICULO 5" / "ARTICLE 5" / "CLÁUSULA 3" / "SECTION 2"
_ARTICLE_HEADER_RE = re.compile(
    r"^(ART[IÍ]CULO|ARTICLE|CL[AÁ]USULA|CLAUSE|SECCI[OÓ]N|SECTION)\s+\d+",
    re.IGNORECASE,
)
#   Bullet / dash list items.
_BULLET_RE = re.compile(r"^[-•*–]\s+(.*)$")
#   Lettered / roman sub-list items: "a)", "(i)", "iii." kept as list items.
_LETTER_LIST_RE = re.compile(r"^(\(?[a-zA-Z]\)|\(?[ivxIVX]+\))\s+\S")

_ALL_CAPS_MAX = 80


def _is_all_caps_title(line: str) -> bool:
    """Short ALL-CAPS line with letters — a section title (preserves the
    pre-existing heuristic so docx_html keeps recognising these)."""
    return 0 < len(line) <= _ALL_CAPS_MAX and line == line.upper() and any(c.isalpha() for c in line)


def _numbered_level(line: str) -> int | None:
    """Heading level (1-3) for a numbered clause, by dotted depth; None if not
    a numbered clause."""
    match = _NUMBERED_CLAUSE_RE.match(line)
    if not match:
        return None
    depth = match.group(1).count(".") + 1
    return min(depth, 3)


def _apply_base_style(document: Document) -> None:
    """Set a sane base font and 1-inch margins on the default section."""
    normal = document.styles["Normal"]
    normal.font.name = _BASE_FONT
    normal.font.size = _BASE_FONT_SIZE
    for section in document.sections:
        section.top_margin = section.bottom_margin = int(_MARGIN_CM * 360000)
        section.left_margin = section.right_margin = int(_MARGIN_CM * 360000)


def _add_flagged_runs(paragraph, text: str) -> None:
    """Add ``text`` to ``paragraph`` as runs, styling any inline markers.

    Every marker is emitted verbatim (so extract_text round-trips it) but in a
    distinct run: defect flags ([MISSING:]/[DEVIATION:]/[REFINEMENT-UNCLEAR:])
    in dark-red bold; verifiable-citation markers ([SOURCE:]) in a small, muted,
    superscript footnote-style run. Surrounding text is a plain run.
    """
    if not any(marker in text for marker in (*_FLAG_MARKERS, _SOURCE_MARKER)):
        paragraph.add_run(text)
        return
    for segment in _FLAG_SPLIT_RE.split(text):
        if not segment:
            continue
        run = paragraph.add_run(segment)
        if segment.startswith(_FLAG_MARKERS):
            run.bold = True
            run.font.color.rgb = _FLAG_HIGHLIGHT
        elif segment.startswith(_SOURCE_MARKER):
            # Footnote-style citation: small, muted, superscript — distinct from
            # the defect highlight, reads as a subtle source reference.
            run.font.size = Pt(8)
            run.font.color.rgb = _SOURCE_FOOTNOTE
            run.font.superscript = True


def _render_body_line(document: Document, line: str) -> None:
    """Map one non-empty body line onto a styled paragraph."""
    # Headings: numbered clauses, ordinal/article headers, ALL-CAPS titles.
    level = _numbered_level(line)
    if level is None and (_ORDINAL_HEADER_RE.match(line) or _ARTICLE_HEADER_RE.match(line)):
        level = 2
    if level is None and _is_all_caps_title(line):
        # ALL-CAPS section titles -> Heading 2 (preserves the long-standing
        # docx_html rendering of these as <h2>).
        level = 2
    if level is not None:
        # Heading styles keep docx_html's OXML walk recognising hierarchy.
        # Flag markers are not expected inside headings, but highlight anyway.
        paragraph = document.add_paragraph(style=f"Heading {level}")
        _add_flagged_runs(paragraph, line)
        return

    # Bullet / dash lists.
    bullet = _BULLET_RE.match(line)
    if bullet:
        paragraph = document.add_paragraph(style="List Bullet")
        _add_flagged_runs(paragraph, bullet.group(1))
        return
    if _LETTER_LIST_RE.match(line):
        paragraph = document.add_paragraph(style="List Bullet")
        _add_flagged_runs(paragraph, line)
        return

    # Ordinary justified body paragraph.
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_flagged_runs(paragraph, line)


def _render_disclaimer(document: Document, disclaimer_lines: list[str]) -> None:
    """Render the SLP disclaimer as a clearly-separated footer block.

    A horizontal divider precedes a small, italic, grey paragraph. The text is
    rendered verbatim (extract_text reads it back unchanged)."""
    divider = document.add_paragraph()
    divider.add_run("—" * 30)  # em-dash rule
    for line in disclaimer_lines:
        if not line.strip():
            continue
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line.strip())
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def render_docx(text: str, title: str | None = None) -> bytes:
    """Convert plain document text into professionally-formatted .docx bytes.

    Recognises numbered clauses, ordinal/article headers and ALL-CAPS titles as
    headings; dash/bullet lines as list items; everything else as justified
    body paragraphs. Flag markers are highlighted but preserved verbatim. The
    SLP disclaimer (if present, already appended to ``text``) is split off and
    rendered once as a footer block.
    """
    document = Document()
    _apply_base_style(document)

    if title:
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(title)
        run.bold = True
        run.font.size = Pt(16)

    # Split off the trailing disclaimer so it can be rendered distinctly. The
    # disclaimer arrives appended verbatim; render the part before it as body
    # and the disclaimer itself as a footer. Anything after the disclaimer
    # (e.g. the Level-3 warning) stays body text.
    body_text = text
    disclaimer_lines: list[str] = []
    idx = text.find(SLP_DISCLAIMER)
    if idx != -1:
        body_text = text[:idx]
        remainder = text[idx + len(SLP_DISCLAIMER):]
        disclaimer_lines = [SLP_DISCLAIMER]
        # Trailing remainder (warnings appended after the disclaimer) renders as
        # body so it is not lost.
        if remainder.strip():
            body_text = body_text + "\n" + remainder

    for line in body_text.split("\n"):
        if not line.strip():
            continue
        _render_body_line(document, line.strip())

    if disclaimer_lines:
        _render_disclaimer(document, disclaimer_lines)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_SECTION_TITLE_MAX = 120


def _is_section_title(line: str) -> bool:
    """Whether a stripped line opens a new section/clause (022 structural
    chunking). Reuses the exact heading heuristics render_docx/docx_html key
    off: ALL-CAPS titles, ARTÍCULO/CLÁUSULA/SECTION headers, Spanish ordinal
    headers and numbered clauses up to depth 2 ("1." / "1.1"). List items are
    never titles."""
    if not line or len(line) > _SECTION_TITLE_MAX:
        return False
    if _BULLET_RE.match(line) or _LETTER_LIST_RE.match(line):
        return False
    if _is_all_caps_title(line):
        return True
    if _ARTICLE_HEADER_RE.match(line) or _ORDINAL_HEADER_RE.match(line):
        return True
    level = _numbered_level(line)
    return level is not None and level <= 2


def split_sections(text: str) -> list[tuple[str | None, str]]:
    """Split plain document text into ``(section_title, section_text)`` blocks.

    Used by the indexer (022) so every RAG chunk records the clause it came
    from — citations become "[1] LPA · CLÁUSULA 8" instead of "[1] LPA".
    The title line is KEPT at the start of its section body (it carries
    meaning: "CLÁUSULA 8 — COMISIONES"). Text before the first title becomes
    an untitled section; a text with no recognizable titles returns one
    untitled section with everything.
    """
    sections: list[tuple[str | None, str]] = []
    title: str | None = None
    buffer: list[str] = []

    def flush(current_title: str | None, lines: list[str]) -> None:
        body = "\n".join(lines).strip()
        if body:
            sections.append((current_title, body))

    for raw in text.split("\n"):
        line = raw.strip()
        if _is_section_title(line):
            flush(title, buffer)
            title = line
            buffer = [raw]
        else:
            buffer.append(raw)
    flush(title, buffer)
    return sections


def extract_text(docx_bytes: bytes) -> str:
    """Extract paragraph text from .docx bytes (used for [MISSING] checks,
    redline bases and RAG context).

    The em-dash divider paragraph inserted before the disclaimer footer carries
    no meaningful text and is dropped so the round-trip stays faithful.
    """
    document = Document(io.BytesIO(docx_bytes))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text
        # Skip the cosmetic divider rule (only em-dashes).
        if text and set(text.strip()) == {"—"}:
            continue
        lines.append(text)
    return "\n".join(lines)
