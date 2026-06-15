"""Redline engine: tracked-changes .docx comparing generated doc vs precedent.

- Paragraph-level alignment with :class:`difflib.SequenceMatcher` aligns
  unchanged / added / removed paragraphs.
- Within a *replaced* block, paragraphs are paired by similarity (best match
  above a ratio threshold) and the matched pair is diffed WORD-by-word, so a
  changed clause shows only the changed words as ``w:ins`` / ``w:del`` rather
  than the whole paragraph being deleted and reinserted. Unpaired old/new
  paragraphs become a deletion / insertion.
- Insertions/deletions are emitted as real Word revision elements
  (``w:ins`` / ``w:del``), author always "Lol-AI-lo AI" (SPEC guardrail 6).
- Pure formatting / whitespace deltas (and, by extension, punctuation-only
  whitespace differences once normalised) are NOT marked.
- **Large documents:** when either side exceeds ``redline_max_paragraphs``
  (config), the expensive similarity pairing + word-level diff is skipped in
  favour of a coarse paragraph-level diff (whole-paragraph ins/del). This is
  logged and still produces a valid tracked-changes .docx — it never crashes
  or hangs on pathological input.
"""
from __future__ import annotations

import difflib
import io
import logging
import re
from datetime import datetime, timezone

from docx import Document
from docx.oxml.ns import qn

from config import get_settings

logger = logging.getLogger("lolailo.redline")

# A replace-block paragraph pair is treated as a *modification* (word-level
# diff) only when its normalised similarity clears this ratio; below it the
# paragraphs are unrelated and emitted as a clean delete + insert.
_PAIR_SIMILARITY_THRESHOLD = 0.4

try:  # python-docx >= 1.x
    from docx.oxml.parser import OxmlElement
except ImportError:  # older layouts
    from docx.oxml import OxmlElement  # type: ignore[attr-defined,no-redef]

REDLINE_AUTHOR = "Lol-AI-lo AI"

_rev_id = 0


def _next_rev_id() -> str:
    global _rev_id
    _rev_id += 1
    return str(_rev_id)


def _now_w3c() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _normalize(text: str) -> str:
    """Collapse whitespace so formatting-only differences compare equal."""
    return re.sub(r"\s+", " ", text).strip()


def _make_run(text: str, del_text: bool = False) -> OxmlElement:
    run = OxmlElement("w:r")
    tag = "w:delText" if del_text else "w:t"
    t = OxmlElement(tag)
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    return run


def _revision_element(kind: str, text: str) -> OxmlElement:
    """Build a w:ins or w:del element wrapping a single run."""
    element = OxmlElement(kind)
    element.set(qn("w:id"), _next_rev_id())
    element.set(qn("w:author"), REDLINE_AUTHOR)
    element.set(qn("w:date"), _now_w3c())
    element.append(_make_run(text, del_text=(kind == "w:del")))
    return element


def _append_plain(paragraph, text: str) -> None:
    if text:
        paragraph._p.append(_make_run(text))


def _append_insertion(paragraph, text: str) -> None:
    if text:
        paragraph._p.append(_revision_element("w:ins", text))


def _append_deletion(paragraph, text: str) -> None:
    if text:
        paragraph._p.append(_revision_element("w:del", text))


def _tokens(text: str) -> list[str]:
    # Keep separators so reassembled text preserves spacing.
    return re.findall(r"\S+\s*", text)


def _diff_paragraph(paragraph, old: str, new: str) -> None:
    """Word-level diff inside a replaced paragraph."""
    old_tokens, new_tokens = _tokens(old), _tokens(new)
    matcher = difflib.SequenceMatcher(a=old_tokens, b=new_tokens, autojunk=False)
    for op, a1, a2, b1, b2 in matcher.get_opcodes():
        if op == "equal":
            _append_plain(paragraph, "".join(old_tokens[a1:a2]))
        else:
            if op in ("replace", "delete"):
                _append_deletion(paragraph, "".join(old_tokens[a1:a2]))
            if op in ("replace", "insert"):
                _append_insertion(paragraph, "".join(new_tokens[b1:b2]))


def _render_replace_block(
    document, old_block: list[str], new_block: list[str], word_level: bool
) -> None:
    """Render a replace block: pair similar paragraphs and word-diff each pair,
    leaving unpaired paragraphs as clean deletions / insertions.

    When ``word_level`` is False (large-document fallback) the block is emitted
    as a coarse whole-paragraph delete-then-insert — no similarity search.
    """
    if not word_level:
        for paragraph_text in old_block:
            _append_deletion(document.add_paragraph(), paragraph_text)
        for paragraph_text in new_block:
            _append_insertion(document.add_paragraph(), paragraph_text)
        return

    # Greedy best-match pairing: for each old paragraph pick the most similar
    # still-unused new paragraph above the threshold. Preserves order for the
    # common case (aligned modified clauses) while tolerating count drift.
    used_new: set[int] = set()
    pairs: dict[int, int] = {}  # old index -> new index
    for oi, old in enumerate(old_block):
        old_norm = _normalize(old)
        best_ratio, best_j = 0.0, -1
        for nj, new in enumerate(new_block):
            if nj in used_new:
                continue
            ratio = difflib.SequenceMatcher(None, old_norm, _normalize(new), autojunk=False).ratio()
            if ratio > best_ratio:
                best_ratio, best_j = ratio, nj
        if best_j != -1 and best_ratio >= _PAIR_SIMILARITY_THRESHOLD:
            pairs[oi] = best_j
            used_new.add(best_j)

    # Emit in new-paragraph order, splicing deletions of unpaired old paragraphs
    # in their original position so the redline reads naturally.
    paired_old_by_new = {nj: oi for oi, nj in pairs.items()}
    next_old = 0

    def flush_deletions_up_to(limit: int) -> None:
        nonlocal next_old
        while next_old < limit:
            if next_old not in pairs:  # old paragraph that found no match
                _append_deletion(document.add_paragraph(), old_block[next_old])
            next_old += 1

    for nj, new in enumerate(new_block):
        if nj in paired_old_by_new:
            oi = paired_old_by_new[nj]
            flush_deletions_up_to(oi)
            _diff_paragraph(document.add_paragraph(), old_block[oi], new)
            next_old = oi + 1
        else:
            _append_insertion(document.add_paragraph(), new)
    flush_deletions_up_to(len(old_block))


def build_redline(precedent_text: str, generated_text: str) -> bytes:
    """Produce a tracked-changes .docx of ``generated_text`` vs ``precedent_text``.

    Aligns paragraphs, word-diffs modified ones, and never marks whitespace-only
    changes. On very large inputs (> ``redline_max_paragraphs`` on either side)
    it falls back to a coarse paragraph-level diff so it can never hang.
    """
    old_paragraphs = [p for p in precedent_text.split("\n") if p.strip()]
    new_paragraphs = [p for p in generated_text.split("\n") if p.strip()]

    # Large-document guard: similarity pairing is O(n*m) per replace block and
    # word-level diff is O(words^2); skip both above the threshold.
    max_paragraphs = get_settings().redline_max_paragraphs
    word_level = max(len(old_paragraphs), len(new_paragraphs)) <= max_paragraphs
    if not word_level:
        logger.warning(
            "Redline input large (%d/%d paragraphs > %d); using coarse "
            "paragraph-level diff.",
            len(old_paragraphs), len(new_paragraphs), max_paragraphs,
        )

    document = Document()
    matcher = difflib.SequenceMatcher(
        a=[_normalize(p) for p in old_paragraphs],
        b=[_normalize(p) for p in new_paragraphs],
        autojunk=False,
    )
    for op, a1, a2, b1, b2 in matcher.get_opcodes():
        if op == "equal":
            # Unchanged (incl. whitespace-only differences): keep the NEW text
            # unmarked — formatting changes are not revisions.
            for paragraph_text in new_paragraphs[b1:b2]:
                _append_plain(document.add_paragraph(), paragraph_text)
        elif op == "delete":
            for paragraph_text in old_paragraphs[a1:a2]:
                _append_deletion(document.add_paragraph(), paragraph_text)
        elif op == "insert":
            for paragraph_text in new_paragraphs[b1:b2]:
                _append_insertion(document.add_paragraph(), paragraph_text)
        else:  # replace
            _render_replace_block(
                document, old_paragraphs[a1:a2], new_paragraphs[b1:b2], word_level
            )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
